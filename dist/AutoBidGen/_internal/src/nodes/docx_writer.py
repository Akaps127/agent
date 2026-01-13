"""
DOCX 공고문 생성 모듈
HTML 출력과 동일한 내용을 DOCX 형식으로 생성합니다.
writer.py의 로직을 그대로 반영하여 4대 파라미터 기반 동적 생성을 지원합니다.
"""
from datetime import datetime
from types import SimpleNamespace
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.schema import PurchasePlan, PlannedNotice
from src.config import get_limit_gosi

# Constants from writer.py
LIMIT_SMALL = 100_000_000

LAW_TEXTS = {
    "국가계약법": {
        "full_name": "국가를 당사자로 하는 계약에 관한 법률",
        "short_name": "국가계약법",
        "시행령": "국가를 당사자로 하는 계약에 관한 법률 시행령",
        "시행규칙": "국가를 당사자로 하는 계약에 관한 법률 시행규칙",
        "부정당업자조항": "제27조",
        "부정당업자조항_금품": "제27조 제1항 제7호",
        "부정당업자조항_담합": "제27조 제1항 제2호",
        "조세포탈조항": "제27조의5",
        "시행령_조세포탈": "제12조제3항",
        "입찰무효조항_시행령": "제39조제4항",
        "입찰무효조항_규칙": "제44조",
        "청렴조항_시행령": "제4조의2 제1항 제2호",
    },
    "지방계약법": {
        "full_name": "지방자치단체를 당사자로 하는 계약에 관한 법률",
        "short_name": "지방계약법",
        "시행령": "지방자치단체를 당사자로 하는 계약에 관한 법률 시행령",
        "시행규칙": "지방자치단체를 당사자로 하는 계약에 관한 법률 시행규칙",
        "부정당업자조항": "제31조",
        "부정당업자조항_금품": "제31조 제1항 제7호",
        "부정당업자조항_담합": "제31조 제1항 제2호",
        "조세포탈조항": "제31조의5",
        "시행령_조세포탈": "제13조제3항",
        "입찰무효조항_시행령": "제42조제4항",
        "입찰무효조항_규칙": "제46조",
        "청렴조항_시행령": "제4조의2 제1항 제2호",
    },
    "자체기준": {
        "full_name": "기관 자체 계약 기준",
        "short_name": "자체기준",
        "시행령": "기관 자체 계약 기준",
        "시행규칙": "기관 자체 계약 기준",
        "부정당업자조항": "해당 조항",
        "부정당업자조항_금품": "해당 조항",
        "부정당업자조항_담합": "해당 조항",
        "조세포탈조항": "해당 조항",
        "시행령_조세포탈": "해당 조항",
        "입찰무효조항_시행령": "해당 조항",
        "입찰무효조항_규칙": "해당 조항",
        "청렴조항_시행령": "해당 조항",
    }
}

CONTRACT_TYPE_DOCS = {
    "공사": "공사입찰유의서",
    "용역": "용역입찰유의서",
    "물품": "물품구매(제조)입찰유의서",
    "물품제조": "물품구매(제조)입찰유의서",
    "물품구매": "물품구매(제조)입찰유의서",
    "외주용역": "용역입찰유의서",
    "학술연구용역": "학술연구용역입찰유의서",
    "시설관리용역": "용역입찰유의서",
    "전문용역": "용역입찰유의서",
    "기술용역": "용역입찰유의서",
}

CONTRACT_TYPE_CONDITIONS = {
    "공사": ("공사계약일반조건", "공사계약특수조건"),
    "용역": ("용역계약일반조건", "용역계약특수조건"),
    "물품": ("물품구매(제조)계약일반조건", "물품구매(제조)계약특수조건"),
    "물품제조": ("물품구매(제조)계약일반조건", "물품구매(제조)계약특수조건"),
    "물품구매": ("물품구매(제조)계약일반조건", "물품구매(제조)계약특수조건"),
    "외주용역": ("용역계약일반조건", "용역계약특수조건"),
    "학술연구용역": ("학술연구용역계약일반조건", "학술연구용역계약특수조건"),
    "시설관리용역": ("용역계약일반조건", "용역계약특수조건"),
    "전문용역": ("용역계약일반조건", "용역계약특수조건"),
    "기술용역": ("용역계약일반조건", "용역계약특수조건"),
}


def set_cell_shading(cell, fill_color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_border_to_paragraph(paragraph, border_width=8, space=4):
    """
    Add border to a paragraph.
    border_width: Width in eighths of a point (8 = 1pt, 16 = 1pt)
    space: Space between border and text in points
    """
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_width))  # 8 = 1pt, 16 = 1pt
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    
    pPr.append(pBdr)


def create_bordered_table(doc, border_width_pt=2):
    """
    Create a single-cell table to simulate a bordered box.
    Returns the table and the cell.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    
    # Set table width to page width (approximately)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 100% width (50 * 100)
    table._element.tblPr.append(tblW)
    
    # Set border style
    tbl_pr = table._element.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    
    border_sz = str(border_width_pt * 8)  # Convert pt to eighths of a point
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tbl_borders.append(border)
    
    tbl_pr.append(tbl_borders)
    
    cell = table.rows[0].cells[0]
    
    # Add padding to cell
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for margin_name in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '150')  # 150 twips ≈ 10.5pt
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    
    tcPr.append(tcMar)
    
    return table, cell


def add_paragraph_with_style(doc, text="", bold=False, blue=False, size=11, 
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, indent_left=0):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    if indent_left > 0:
        p.paragraph_format.left_indent = Inches(indent_left)
    
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        if bold:
            run.bold = True
        if blue:
            run.font.color.rgb = RGBColor(0, 0, 255)
    
    return p


def add_mixed_paragraph(doc, parts, indent_left=0):
    """
    Add a paragraph with mixed formatting.
    parts: list of tuples (text, bold, blue)
    """
    p = doc.add_paragraph()
    if indent_left > 0:
        p.paragraph_format.left_indent = Inches(indent_left)
    
    for text, bold, blue in parts:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        if bold:
            run.bold = True
        if blue:
            run.font.color.rgb = RGBColor(0, 0, 255)
    
    return p


def generate_docx(purchase_plan: PurchasePlan, planned_notice: PlannedNotice, output_path: str = None) -> str:
    """
    Generate a DOCX document from the planned notice data.
    Mirrors the logic in writer.py for consistency.
    
    Args:
        purchase_plan: The original purchase plan data
        planned_notice: The processed notice data from planner
        output_path: Optional output path. If None, uses default.
        
    Returns:
        Path to the generated DOCX file
    """
    if output_path is None:
        output_path = f"output_notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # --- 0. 4대 파라미터 결정 (from writer.py) ---
    contract_law = planned_notice.contract_law_type or getattr(purchase_plan, 'contract_law_type', None) or "국가계약법"
    contract_type = planned_notice.contract_type or getattr(purchase_plan, 'contract_type', None) or "물품구매"
    winner_method = planned_notice.winner_determination or getattr(purchase_plan, 'winner_determination', None) or planned_notice.notice_type
    
    if winner_method == "소액수의":
        bidding_method = "수의계약"
    else:
        bidding_method = planned_notice.bidding_method or getattr(purchase_plan, 'bidding_method', None) or "제한경쟁"
    
    # 법령 텍스트 가져오기
    law = LAW_TEXTS.get(contract_law, LAW_TEXTS["국가계약법"])
    bid_doc = CONTRACT_TYPE_DOCS.get(contract_type, "물품구매(제조)입찰유의서")
    general_cond, special_cond = CONTRACT_TYPE_CONDITIONS.get(contract_type, 
                                                              ("물품구매(제조)계약일반조건", "물품구매(제조)계약특수조건"))
    
    # --- 1. Data Preparation (from writer.py) ---
    data = SimpleNamespace()
    data.winning_method = winner_method
    data.bidding_method = bidding_method
    data.item_codes = purchase_plan.item_codes
    data.industry_info = getattr(purchase_plan, 'industry_info', [])
    
    item_names_raw = getattr(purchase_plan, 'item_names', [])
    if isinstance(item_names_raw, list):
        data.item_names = ", ".join(item_names_raw) if item_names_raw else "물품"
    else:
        data.item_names = item_names_raw or "물품"
    
    data.company_restriction = purchase_plan.sme_restriction_text
    data.joint_contract_allow = purchase_plan.joint_venture_allow
    
    winning_rate = getattr(purchase_plan, 'winning_rate', None)
    if winning_rate:
        data.winning_rate = winning_rate
    else:
        data.winning_rate = "88" if data.winning_method == "소액수의" else "84.245"
    
    data.delivery_term = purchase_plan.delivery_period_text
    
    # Contract method sentence
    if data.winning_method == "소액수의":
        data.contract_method_sentence = "소액수의(총액, 전자) 대상입니다."
    elif bidding_method == "제한경쟁":
        data.contract_method_sentence = "제한경쟁(총액), 전자입찰대상 물품입니다."
    elif bidding_method == "일반경쟁":
        data.contract_method_sentence = "일반경쟁(총액), 전자입찰대상 물품입니다."
    elif bidding_method == "지명경쟁":
        data.contract_method_sentence = "지명경쟁(총액), 전자입찰대상 물품입니다."
    else:
        data.contract_method_sentence = f"{bidding_method}(총액), 전자입찰대상 물품입니다."
    
    # 적격심사 관련 문구
    if data.winning_method == "소액수의":
        data.qualification_ref = "제외대상입니다."
    else:
        q_val = planned_notice.qualification_sentences[0] if planned_notice.qualification_sentences else ""
        data.qualification_ref = f"대상 물품입니다. {q_val} 적용"
    
    # Contact Info Helper
    def parse_contact(contact):
        info = {
            "dept": "본사 및 소속기관 사업부서",
            "name": "담당자",
            "tel": ""
        }
        
        if not contact:
            return info
        
        if isinstance(contact, str):
            info["name"] = contact
            return info
        
        if isinstance(contact, dict):
            if contact.get("department"): info["dept"] = contact.get("department")
            if contact.get("name"): info["name"] = contact.get("name")
            if contact.get("phone"): info["tel"] = contact.get("phone")
        else:
            if hasattr(contact, "department") and contact.department: info["dept"] = contact.department
            if hasattr(contact, "name") and contact.name: info["name"] = contact.name
            if hasattr(contact, "phone") and contact.phone: info["tel"] = contact.phone
        
        return info
    
    project_info = parse_contact(purchase_plan.project_contact)
    data.project_contact_dept = project_info["dept"]
    data.project_contact_name = project_info["name"]
    data.project_contact_tel = project_info["tel"]
    
    contract_info = parse_contact(planned_notice.contract_contact)
    data.contract_contact_dept = contract_info["dept"] or "경영지원처 계약부"
    data.contract_contact_name = contract_info["name"]
    data.contract_contact_tel = contract_info["tel"]
    
    # Other variables
    raw_name = purchase_plan.notice_name or ""
    clean_notice_name = raw_name.replace("계획(안)", "").replace("계획안", "").strip()
    formatted_budget = f"{purchase_plan.budget_total:,}"
    
    # Date Formatting
    def format_g2b_date(date_str):
        if not date_str: return ""
        try:
            dt = datetime.strptime(date_str.strip(), "%Y-%m-%d %H:%M")
            return dt.strftime("%Y. %m. %d.(%H:%M)")
        except ValueError:
            return date_str
    
    submission_period_text = ""
    opening_date_text = ""
    
    if purchase_plan.bid_submission_start and purchase_plan.bid_submission_end:
        start_str = format_g2b_date(purchase_plan.bid_submission_start)
        end_str = format_g2b_date(purchase_plan.bid_submission_end)
        
        if len(start_str) > 4 and len(end_str) > 4 and start_str[:4] == end_str[:4]:
            end_str_short = end_str[6:]
            submission_period_text = f"{start_str} ~ {end_str_short}"
        else:
            submission_period_text = f"{start_str} ~ {end_str}"
    else:
        submission_period_text = "공고서 참조"
    
    if purchase_plan.bid_opening_datetime:
        opening_date_text = format_g2b_date(purchase_plan.bid_opening_datetime)
    else:
        opening_date_text = "입찰서 제출 마감일 직후"
    
    data.submission_period_text = submission_period_text
    data.opening_date_text = opening_date_text
    
    sec1_tel = data.project_contact_tel or "000-0000-0000"
    
    def format_sec10(dept, name, tel):
        parts = []
        if dept: parts.append(dept)
        if name: parts.append(name)
        base = " ".join(parts)
        if tel:
            return f"{base} (☎{tel})"
        return base
    
    sec10_project = format_sec10(data.project_contact_dept, data.project_contact_name, data.project_contact_tel)
    sec10_contract = format_sec10(data.contract_contact_dept, data.contract_contact_name, data.contract_contact_tel)
    
    today_str = datetime.now().strftime("%Y년 %m월 %d일")
    
    # Title - 계약유형 및 낙찰자결정방법에 따라 변경
    if data.winning_method == "소액수의":
        title = "소액수의 견적제출 공고"
    else:
        if contract_type in ["공사"]:
            title = "공사 입찰공고"
        elif contract_type in ["용역", "외주용역", "학술연구용역", "시설관리용역", "전문용역", "기술용역"]:
            title = "용역 입찰공고"
        else:
            title = "물품구매 입찰공고"
    
    # Budget-based restriction logic
    budget_supply = purchase_plan.budget_supply
    
    # --- Create Document ---
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    # Notice number
    notice_num_p = doc.add_paragraph()
    notice_num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = notice_num_p.add_run("한국환경공단 입찰공고번호 : 00-00000000-00")
    run.font.size = Pt(11)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    doc.add_paragraph()
    
    # Integrity Box - Using bordered table to match HTML
    integrity_table, integrity_cell = create_bordered_table(doc, border_width_pt=2)
    
    # Title inside box
    title_p = integrity_cell.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("< 본 계약은 청렴계약제가 적용됩니다 >")
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    # First paragraph
    p1 = integrity_cell.add_paragraph()
    run = p1.add_run(f"이 계약은 「{law['full_name']}」 또는 「지방자치단체를 당사자로 하는 계약에 관한 법률」에 따른 청렴계약제가 적용됩니다. 입찰자는 반드시 입찰서 제출 시 아래 청렴계약서에 관한 내용을 숙지·승낙하여야 하며, 동 내용을 위반한 경우 발주기관의 조치에 대해서 어떠한 이의도 제기할 수 없습니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p1.paragraph_format.space_after = Pt(10)
    
    # Main integrity text
    p2 = integrity_cell.add_paragraph()
    run = p2.add_run(f"「{law['full_name']}」 또는 「지방자치단체를 당사자로 하는 계약에 관한 법률」에 따라 본 입찰에 참여한 당사 대리인과 임직원은 입찰·낙찰, 계약 체결 및 이행, 감독, 검사 등의 과정(준공·납품 이후를 포함한다)에서 아래 각호의 청렴계약 조건을 준수할 것이며, 이를 위반한 때에는 입찰·낙찰을 취소하거나 계약을 해제·해지하는 등의 불이익을 감수하고, 이에 민·형사상 이의를 제기하지 않을 것임을 약정합니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p2.paragraph_format.space_after = Pt(10)
    
    # Item 1
    p3 = integrity_cell.add_paragraph()
    run = p3.add_run(f"1. 금품·향응 등(친인척 등에 대한 부정한 취업 제공 포함)을 요구 또는 약속하거나 수수(授受)하지 않을 것이며, 관계공무원에게 금품, 향응 등을 제공한 경우에는 「{law['full_name']}」{law['부정당업자조항_금품']} 또는 「지방자치단체를 당사자로 하는 계약에 관한 법률」제31조 제1항 제7호에 따른 부정당업자의 입찰참가자격 제한 처분을 받겠습니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p3.paragraph_format.space_after = Pt(10)
    
    # Item 2
    p4 = integrity_cell.add_paragraph()
    run = p4.add_run(f"2. 입찰가격의 사전 협의 또는 특정인의 낙찰을 위한 담합 등 공정한 경쟁을 방해하는 행위시에는 「{law['full_name']}」{law['부정당업자조항_담합']} 또는 「지방자치단체를 당사자로 하는 계약에 관한 법률」 제31조 제1항 제2호에 따른 부정당업자 입찰참가자격 제한 처분을 받겠습니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p4.paragraph_format.space_after = Pt(10)
    
    # Inner box (담합 관련) - Using nested table with 1pt border
    inner_table, inner_cell = create_bordered_table(doc, border_width_pt=1)
    # Remove the table from doc and add to integrity_cell instead
    doc._element.body.remove(inner_table._element)
    integrity_cell._element.append(inner_table._element)
    
    inner_p = inner_cell.add_paragraph()
    run = inner_p.add_run("우리 공단은 입찰담합 방지 및 공정거래질서 확립을 위해 「독점규제 및 공정거래에 관한 법률」에 따라 입찰담합징후분석시스템에 입찰정보를 제공하고 있습니다. 입찰담합징후 발견 시 공정거래위원회 제보 및 경찰 조사의뢰 등을 검토·시행하고 있으며, 입찰담합으로 판명시 부정당업자 제재(입찰참가자격제한) 처분 및 손해배상청구소송 제소 등 법적 제재조치를 시행하고 있습니다.")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    # Item 3
    p5 = integrity_cell.add_paragraph()
    run = p5.add_run("3. 공정한 직무수행을 방해하는 알선·청탁을 통하여 입찰 또는 계약과 관련된 특정 정보의 제공을 요구하거나 받는 행위를 하지 않겠습니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p5.paragraph_format.space_after = Pt(10)
    p5.paragraph_format.space_before = Pt(10)
    
    # Item 4
    p6 = integrity_cell.add_paragraph()
    run = p6.add_run(f"4. 「{law['시행령']}」 {law['청렴조항_시행령']} 위반 시에 아래의 손해배상액을 납부토록 하겠습니다.")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    p7 = integrity_cell.add_paragraph()
    run = p7.add_run("  - 입찰자 : 입찰금액의 100분의 5")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    p8 = integrity_cell.add_paragraph()
    run = p8.add_run("  - 계약상대자 : 계약금액의 100분의 10")
    run.font.size = Pt(10.5)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    doc.add_paragraph()
    
    # Section 1: 견적(입찰)에 부치는 사항
    add_paragraph_with_style(doc, "1. 견적(입찰)에 부치는 사항", bold=True, size=12)
    
    add_mixed_paragraph(doc, [("가. 공고명 : ", False, False), (clean_notice_name, True, True)], indent_left=0.25)
    add_mixed_paragraph(doc, [("나. 계약기간 : ", False, False), (data.delivery_term, True, True)], indent_left=0.25)
    add_mixed_paragraph(doc, [("다. 예산액 : ", False, False), (f"{formatted_budget}원", True, True), ("(부가가치세 포함)", False, False)], indent_left=0.25)
    add_mixed_paragraph(doc, [("라. 구매범위 : 물품규격서 등 참조(문의 ☎", False, False), (sec1_tel, True, True), (", ", False, False), (data.project_contact_name, True, True), (")", False, False)], indent_left=0.25)
    add_mixed_paragraph(doc, [("마. 전자입찰서 제출기간 : ", False, False), (data.submission_period_text, True, True)], indent_left=0.25)
    add_mixed_paragraph(doc, [("바. 개찰일시 및 장소 : ", False, False), (data.opening_date_text, True, True), (", 국가종합전자조달시스템(나라장터)", False, False)], indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 2: 견적(입찰) 및 계약방식
    add_paragraph_with_style(doc, "2. 견적(입찰) 및 계약방식", bold=True, size=12)
    
    add_mixed_paragraph(doc, [("가. ", False, False), (data.contract_method_sentence, True, True)], indent_left=0.25)
    add_mixed_paragraph(doc, [("나. ", False, False), (f"적격심사 {data.qualification_ref}", True, True)], indent_left=0.25)
    add_paragraph_with_style(doc, "다. 청렴계약이행 서약제 대상입니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "라. 입찰서는 반드시 국가종합전자조달시스템(www.g2b.go.kr)의 전자입찰특별유의서에 따라 제출하여야 합니다.", size=11, indent_left=0.25)
    
    add_paragraph_with_style(doc, "※ 입찰 전 납품규격, 납품가능 금액, 납품가능 여부 등을 반드시 확인하시기를 바라며, 이에 대한 검토 없이 무리하게 저가 입찰한 책임은 입찰참가자에게 있음을 알려드립니다.", size=10, indent_left=0.5)
    add_paragraph_with_style(doc, "※ 기타 세부사항은 전자입찰 공고서에 첨부된 규격서(시방서), 과업내용서 등을 반드시 확인하신 후 과업이행에 필요한 총금액을 산출하여 투찰하시기 바랍니다.", size=10, indent_left=0.5)
    
    add_paragraph_with_style(doc, "마. 입찰금액은 반드시 부가가치세를 포함한 금액으로 제출하여야 하며 비영리법인 등 부가가치세 면제대상인 경우 견적금액에서 부가가치세를 차감한 금액을 계약금액으로 결정합니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "바. 정부입찰·계약집행기준 제10조의2 제2항제7호에 따라 전자입찰서 제출 후 정당한 이유없이 계약에 응하지 아니하거나 포기서를 제출하는 경우에는 나라장터 전자조달시스템에 수의계약배제업체로 등록되며, 등록일로부터 3개월간 공단과의 소액수의 계약이 제한됩니다.", size=11, indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 3: 입찰참가자격
    add_paragraph_with_style(doc, "3. 입찰참가자격 : 아래의 입찰참가자격을 모두 갖춘 자이어야 합니다.", bold=True, size=12)
    
    # Item codes
    if data.item_codes:
        add_paragraph_with_style(doc, "가. 국가종합전자조달시스템 입찰참가자격등록규정에 따라 반드시 전자입찰서 제출 마감일 전일까지 나라장터(G2B)시스템에 아래의 사항을 모두 입찰참가자격으로 등록한 자", size=11, indent_left=0.25)
        
        code = data.item_codes[0] if isinstance(data.item_codes, list) else data.item_codes
        i_name = "물품"
        if isinstance(data.item_names, list) and data.item_names:
            i_name = data.item_names[0]
        elif isinstance(data.item_names, str):
            i_name = data.item_names
        
        add_mixed_paragraph(doc, [("○ 입찰참가 등록 마감일 기준 ", False, False), (i_name, True, True), 
                                  ("(세부품명번호 10자리 ", False, False), (code, True, True), 
                                  (")를 제조 또는 공급 물품으로 입찰참가 등록한 자", False, False)], indent_left=0.5)
    else:
        add_paragraph_with_style(doc, "가. 국가종합전자조달시스템 입찰참가자격등록규정에 따라 반드시 전자입찰서 제출 마감일 전일까지 나라장터(G2B)시스템에 아래의 사항을 모두 입찰참가자격으로 등록한 자", size=11, indent_left=0.25)
        add_paragraph_with_style(doc, "○ 세부품명번호(10자리)를 제조 또는 공급 물품으로 입찰참가 등록한 자", size=11, indent_left=0.5)
    
    # Industry codes
    if hasattr(data, 'industry_info') and data.industry_info:
        info = data.industry_info[0]
        code = info.industry_code
        name = info.industry_name
        law_name = info.legal_bases[0]['name'] if info.legal_bases and len(info.legal_bases) > 0 else "관련 법령"
        
        if name:
            add_mixed_paragraph(doc, [("○ 「", False, False), (law_name, True, True), ("」에 의한 ", False, False), 
                                      (name, True, True), ("(업종코드: ", False, False), (code, True, True), 
                                      (")으로 입찰참가 등록한 자", False, False)], indent_left=0.5)
        else:
            add_mixed_paragraph(doc, [("○ 관련 법령에 따른 업종코드 ", False, False), (code, True, True), 
                                      ("를 등록한 자", False, False)], indent_left=0.5)
    elif hasattr(purchase_plan, 'industry_names') and purchase_plan.industry_names and purchase_plan.industry_codes:
        code = purchase_plan.industry_codes[0] if isinstance(purchase_plan.industry_codes, list) else purchase_plan.industry_codes
        ind_name = purchase_plan.industry_names[0] if isinstance(purchase_plan.industry_names, list) else purchase_plan.industry_names
        
        add_mixed_paragraph(doc, [("○ 관련 법령에 따른 ", False, False), (ind_name, True, True), 
                                  ("(업종코드: ", False, False), (code, True, True), 
                                  (")를 등록한 자", False, False)], indent_left=0.5)
    elif hasattr(purchase_plan, 'industry_codes') and purchase_plan.industry_codes:
        code = purchase_plan.industry_codes[0] if isinstance(purchase_plan.industry_codes, list) else purchase_plan.industry_codes
        add_mixed_paragraph(doc, [("○ 관련 법령에 따른 업종코드 ", False, False), (code, True, True), 
                                  ("를 등록한 자", False, False)], indent_left=0.5)
    else:
        add_paragraph_with_style(doc, "○ 관련 법령에 따른 업종등록을 필한 자", size=11, indent_left=0.5)
    
    # 부정당업자 조항
    add_paragraph_with_style(doc, f"나. 「{law['full_name']}」 {law['부정당업자조항']}(부정당업자의 입찰참가 자격제한)에 해당되지 아니한 업체", size=11, indent_left=0.25)
    
    # 조세포탈 조항
    tax_text = f"""다. 「{law["full_name"]}」 {law["조세포탈조항"]} 및 같은 법 시행령 {law["시행령_조세포탈"]}에 따라 '조세포탈 등을 한 자'로서 유죄판결이 확정된 날부터 2년이 지나지 아니한 자는 입찰에 참여할 수 없습니다.
입찰자는 같은 법 시행령 {law["시행령_조세포탈"]} 각 호에 해당하지 아니한다는 서약서를 입찰시 제출하여야 합니다. 만일 서약내용이 허위로 판명될 경우 계약의 해제·해지를 당할 수 있고, 부정당업자 입찰참가자격제한처분을 받을 수 있습니다.
다만, 나라장터 시스템을 이용하여 제출하는 경우에는 전자입찰서에 동 서약서의 내용을 포함하고 있으므로 전자입찰서 제출로 서약서 제출을 갈음합니다."""
    add_paragraph_with_style(doc, tax_text, size=11, indent_left=0.25)
    
    # Restriction details based on budget
    if budget_supply < LIMIT_SMALL:
        # 소기업
        restriction_text = """라. 소기업·소상공인 제한 경쟁입니다.
  ○ 소기업 제한 : 「중소기업기본법」 제2조에 따른 소기업 또는 「소상공인 보호 및 지원에 관한 법률」 제2조에 따른 소상공인으로서 「중소기업 범위 및 확인에 관한 규정」에 따라 발급된 소기업·소상공인확인서 (전자입찰서 제출마감일 전일까지 발급된 것으로 유효기간내 있어야 함)를 소지한 업체이어야 합니다.
  ※ 「중소기업제품 구매촉진 및 판로지원에 관한 법률」 제33조 제1항에 따라 소기업으로 간주되는 특별법인으로 중소기업제품 「공공구매제도 운영요령」 제45조에 따라 '특별법인 소기업 간주확인서'를 소기업 또는 소상공인으로 발급받은 경우 입찰참가 자격이 있으며 상기 각 호의 입찰참가자격을 모두 갖추어야 합니다.
  ※ <중소기업·소상공인확인서>는 중소기업공공구매 종합정보망에서 확인하며 확인되지 않을 경우 입찰참가자격이 없습니다."""
        add_mixed_paragraph(doc, [("라. ", False, False), ("소기업·소상공인", True, True), (" 제한 경쟁입니다.", False, False)], indent_left=0.25)
        add_paragraph_with_style(doc, "  ○ 소기업 제한 : 「중소기업기본법」 제2조에 따른 소기업 또는 「소상공인 보호 및 지원에 관한 법률」 제2조에 따른 소상공인으로서 「중소기업 범위 및 확인에 관한 규정」에 따라 발급된 소기업·소상공인확인서 (전자입찰서 제출마감일 전일까지 발급된 것으로 유효기간내 있어야 함)를 소지한 업체이어야 합니다.", size=11, indent_left=0.5)
        add_paragraph_with_style(doc, "  ※ 「중소기업제품 구매촉진 및 판로지원에 관한 법률」 제33조 제1항에 따라 소기업으로 간주되는 특별법인으로 중소기업제품 「공공구매제도 운영요령」 제45조에 따라 '특별법인 소기업 간주확인서'를 소기업 또는 소상공인으로 발급받은 경우 입찰참가 자격이 있으며 상기 각 호의 입찰참가자격을 모두 갖추어야 합니다.", size=11, indent_left=0.5)
        add_paragraph_with_style(doc, "  ※ <중소기업·소상공인확인서>는 중소기업공공구매 종합정보망에서 확인하며 확인되지 않을 경우 입찰참가자격이 없습니다.", size=11, indent_left=0.5)
    elif budget_supply < get_limit_gosi():
        # 중소기업
        add_mixed_paragraph(doc, [("라. ", False, False), ("중소기업·소상공인", True, True), (" 제한 경쟁입니다.", False, False)], indent_left=0.25)
        add_paragraph_with_style(doc, "  ○ 중소기업 제한 : 「중소기업기본법」 제2조에 따른 중소기업 또는 「소상공인 보호 및 지원에 관한 법률」 제2조에 따른 소상공인으로서 「중소기업 범위 및 확인에 관한 규정」에 따라 발급된 중소기업·소상공인확인서 (전자입찰서 제출마감일 전일까지 발급된 것으로 유효기간내 있어야 함)를 소지한 업체이어야 합니다.", size=11, indent_left=0.5)
        add_paragraph_with_style(doc, "  ※ 「중소기업제품 구매촉진 및 판로지원에 관한 법률」 제33조 제1항에 따라 소기업으로 간주되는 특별법인으로 중소기업제품 「공공구매제도 운영요령」 제45조에 따라 '특별법인 소기업 간주확인서'를 소기업 또는 소상공인으로 발급받은 경우 입찰참가 자격이 있으며 상기 각 호의 입찰참가자격을 모두 갖추어야 합니다.", size=11, indent_left=0.5)
        add_paragraph_with_style(doc, "  ※ <중소기업·소상공인확인서>는 중소기업공공구매 종합정보망에서 확인하며 확인되지 않을 경우 입찰참가자격이 없습니다.", size=11, indent_left=0.5)
    else:
        # 제한 없음
        add_paragraph_with_style(doc, "라. 본 입찰은 기업구분에 따른 입찰참가 제한이 없습니다.", size=11, indent_left=0.25)

    if "달력" in data.item_names:
        add_paragraph_with_style(doc, "  ○ 「중소기업제품 구매촉진 및 판로지원에 관한 법률」 제9조 및 동법 시행규칙 제5조 규정에 의한 직접생산확인증명서[세부품명번호 : 달력(4411200201)]를 소지한 자(개찰일 전일까지 발급된 것으로 유효기간 내에 있어야 함)", size=11, indent_left=0.5)
    
    doc.add_paragraph()
    
    # Section 4: 공동계약
    add_paragraph_with_style(doc, "4. 공동계약", bold=True, size=12)
    joint_text = '본 계약은 공동수급을 허용합니다.' if data.joint_contract_allow else '본 계약은 공동수급을 허용하지 않습니다.'
    add_mixed_paragraph(doc, [(joint_text, True, True)], indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 5: 예정가격 및 낙찰자 결정방법
    add_paragraph_with_style(doc, "5. 예정가격 및 낙찰자 결정방법", bold=True, size=12)
    add_paragraph_with_style(doc, "가. 예정가격은 예비가격기초금액기준 ±2% 범위내에서 작성된 15개 복수 예비가격 중 입찰에 참여하는 각 업체가 추첨(2개씩 선택)한 번호 중 가장 많이 선택된 4개의 예비가격을 산술평균한 가격으로 결정됩니다.", size=11, indent_left=0.25)
    add_mixed_paragraph(doc, [("나. 낙찰자(계약상대자) 선정은 예정가격의 ", False, False), (f"{data.winning_rate}%", True, True), 
                              (" 이상으로 견적서를 제출한 자 중 최저가격으로 견적서를 제출한 자 순서에 따라 「공직자의 이해충돌방지법」 제12조제1항 수의계약 체결 제한 사유에 해당하지 아니한 자를 계약상대자로 결정합니다.", False, False)], indent_left=0.25)
    add_paragraph_with_style(doc, f"다. 낙찰이 될 수 있는 동일가격으로 견적 제출한 자가 2인 이상일 때에는 {law['short_name']} 시행령 제47조 규정에 의거 낙찰자를 결정합니다. 전자입찰유의서 제15조에 따라 추첨에 의해 낙찰자를 결정하는 경우 전자조달시스템을 통한 자동추첨방식을 적용하여 계약상대자를 결정합니다.", size=11, indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 6: 청렴계약이행 서약서 제출
    add_paragraph_with_style(doc, "6. 청렴계약이행 서약서 제출", bold=True, size=12)
    add_paragraph_with_style(doc, "가. 입찰에 참여한 자는 모두 청렴계약이행을 위한 공정경쟁 및 청렴계약 입찰특별유의서 제3조에 의거 청렴계약이행서약서를 제출한 것으로 갈음합니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "· 우리공단은 청렴계약 실효성 확보를 위한 입찰담합방지책으로 손해배상제도를 시행하고 있으니 유의하여 주시기 바랍니다.", size=11, indent_left=0.5)
    add_paragraph_with_style(doc, "· 관련자료는 한국환경공단 홈페이지(www.keco.or.kr) 입찰정보/집행기준에서 열람 및 다운 받을 수 있습니다.", size=11, indent_left=0.5)
    
    doc.add_paragraph()
    
    # Section 7: 입찰보증금 납부 및 동 귀속
    add_paragraph_with_style(doc, "7. 입찰보증금 납부 및 동 귀속", bold=True, size=12)
    add_paragraph_with_style(doc, "가. 소액수의계약은 경쟁입찰이 아니므로 입찰보증금은 납부받지 아니합니다.", size=11, indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 8: 입찰무효 또는 취소
    add_paragraph_with_style(doc, "8. 입찰무효 또는 취소", bold=True, size=12)
    add_paragraph_with_style(doc, f"가. 「{law['시행령']}」 {law['입찰무효조항_시행령']}, 같은 법 시행규칙 {law['입찰무효조항_규칙']} 및 「(계약예규){bid_doc}」 제12조에 해당되는 입찰은 무효입니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "나. 입찰참가자격등록증상의 상호 및 대표자(수인대표인 경우 대표자 전원의 성명을 모두 등재, 각자 대표도 해당)가 법인등기부등본상의 상호, 대표자와 다른 경우에는 입찰참가자격등록증을 변경등록하고 입찰에 참여하여야 하며, 변경등록하지 않고 참여한 입찰은 무효입찰임을 알려드립니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "다. 입찰참가자격의 판단기준일은 입찰참가자격등록 마감일(기준일이 정해져 있는 경우에는 해당일)이며 마감일까지 참가자격을 갖추지 않은 경우 무효입찰입니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, f"라. 「{law['시행규칙']}」 {law['입찰무효조항_규칙']} 및 「(계약예규){bid_doc}」 제12조에 정한 입찰무효 해당 여부 확인을 위하여 등록정보 확인을 위한 서류(법인등기부등본, 입찰대리인임을 증명하는 서류, 개인정보수집이용동의서 등)를 요청하는 경우, 낙찰대상자는 관계 서류를 제출하여야 합니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "마. 전자입찰의 취소 신청은 「국가종합전자조달시스템 전자입찰특별유의서」에 따라 전자입찰서 제출 마감시간 전까지 하셔야 하며, 취소 시 재입찰을 할 수 없습니다.", size=11, indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 9: 하도급에 관한 사항
    add_paragraph_with_style(doc, "9. 하도급에 관한 사항", bold=True, size=12)
    add_paragraph_with_style(doc, "가. 본 계약은 하도급 불가 건으로 개별 법령상 하도급 규정을 위반하여 하도급을 하거나, 발주기관 승인 없이 하도급을 하는 경우 부정당업자로 입찰참가자격 제한을 받을 수 있습니다.", size=11, indent_left=0.25)
    
    doc.add_paragraph()
    
    # Section 10: 기타사항 및 추가정보 제공처
    add_paragraph_with_style(doc, "10. 기타사항 및 추가정보 제공처", bold=True, size=12)
    add_paragraph_with_style(doc, f"가. 입찰에 참여하고자 하는 자는 공고서, 규격서, 과업내용서, 적격심사기준(적격심사 대상물품에 한함), {general_cond}, {special_cond}, {bid_doc}, 국가종합전자조달시스템 전자입찰특별유의서, 청렴계약 입찰특변유의서 및 이행각서 등 입찰에 필요한 모든 사항을 완전히 숙지하고 입찰에 참여하여야 하며, 이를 숙지하지 못하여 발생하는 책임은 입찰자에게 있습니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "나. 규격 착오 또는 규정의 미숙지 등으로 입찰자가 계약을 체결하지 않거나, 계약을 체결하고 불이행하는 경우 관계 법령에 따라 부정당업자로 제재되어 일정기간 입찰참여가 제한되는 등 불이익을 받으실 수 있으니, 본 입찰공고서의 규격서 및 계약관련 규정을 철저히 숙지하신 후 입찰에 참가하시기 바랍니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "다. 본 입찰은 국가종합전자조달시스템 전자입찰 특별유의서 제7조에 따른 신원확인 입찰이 적용되며, 개인인증서를 보유한 대표자 또는 입찰대리인은 국가종합전자조달시스템전자입찰특별유의서 제7조 제1항 제5호에 따라 미리 지문정보를 등록하여야 전자입찰서 제출이 가능합니다. 다만, 지문인식신원확인 입찰이 곤란한 자는 국가종합전자조달시스템 전자입찰특별유의서 제7조 제1항 제6호 및 제7호의 절차에 따라 예외적으로 개인인증서에 의한 전자입찰서 제출이 가능합니다.", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "라. 기타 문의사항", size=11, indent_left=0.25)
    add_paragraph_with_style(doc, "○ 전자입찰이용안내 : 국가종합전자조달시스템 콜센터(☎1588-0800)", size=11, indent_left=0.5)
    add_mixed_paragraph(doc, [("○ 물품 규격 등 관련사항 : ", False, False), (sec10_project, True, True)], indent_left=0.5)
    add_mixed_paragraph(doc, [("○ 입찰 계약 관련사항 : ", False, False), (sec10_contract, True, True)], indent_left=0.5)
    
    doc.add_paragraph()
    
    # 이의제기 및 신고채널 안내 - Using bordered table to match HTML
    complaint_table, complaint_cell = create_bordered_table(doc, border_width_pt=2)
    
    # Title inside box
    title_p = complaint_cell.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("< 이의제기 및 신고채널 안내 >")
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    # Content
    p1 = complaint_cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("◎ 본 입찰과 관련한 부당행위 또는 부당사례 등과 공단 직원이 금품 및 향응요구, 지위남용 등 부당한 요구를 할 경우 아래 신고채널을 통해 신고할 수 있으며, 신고에 따른 일체의 불이익은 없습니다.")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p1.paragraph_format.space_after = Pt(10)
    
    p2 = complaint_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("K-eco 신문고 공단홈페이지(www.keco.or.kr) > 국민참여 > K-eco신문고")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    p3 = complaint_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("- 부패신고센터 : 전화 032-590-3072 FAX 032-590-3069")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    add_paragraph_with_style(doc, "위와 같이 공고합니다.", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_mixed_paragraph(doc, [(today_str, True, True)], indent_left=0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_paragraph_with_style(doc, "한국환경공단 계약담당", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Save document
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # Test
    from src.schema import PurchasePlan, PlannedNotice
    
    mock_plan = PurchasePlan(
        notice_name="수질자동측정망 측정장비 구매",
        budget_total=60000000,
        budget_supply=54545454,
        contract_method_text="소액수의",
        delivery_period_text="계약일로부터 120일",
        project_contact="032-123-4567",
        contract_contact="032-987-6543"
    )
    mock_planned = PlannedNotice(
        notice_type="소액수의",
        sme_restriction="소기업",
        submission_period="게시일로부터 3일간",
        contract_method_sentence="소액수의(견적입찰)",
        qualification_sentences=["소기업 확인서 소지자"],
        joint_contract_sentence="공동계약 불허",
        notice_name="수질자동측정망 측정장비 구매",
        budget_format="60,000,000원",
        project_contact="032-123-4567",
        contract_contact="032-987-6543",
        bid_submission_start="2025-12-19 09:00",
        bid_submission_end="2025-12-24 10:00",
        bid_opening_datetime="2025-12-24 11:00",
        bid_opening_place="국가종합전자조달시스템(나라장터)"
    )
    
    path = generate_docx(mock_plan, mock_planned, "test_output.docx")
    print(f"Generated: {path}")
